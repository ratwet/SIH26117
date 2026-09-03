"""
SovereignWorkbench — Word Approval Note Compiler (app/compilers/docx_builder.py)
Generates high-fidelity, corporate executive evaluation and approval dossiers
for Mangalore Refinery and Petrochemicals Limited (MRPL) using python-docx.
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

from app.schemas import ApprovalNotePayload


# Corporate Colors
MRPL_NAVY = RGBColor(0x12, 0x34, 0x56)     # #123456
ACCENT_BLUE = RGBColor(0x1B, 0x36, 0x5D)   # #1B365D
CRITICAL_RED = RGBColor(0xC0, 0x00, 0x00)  # #C00000
TEXT_DARK = RGBColor(0x22, 0x22, 0x22)     # #222222
MUTED_GRAY = RGBColor(0x55, 0x55, 0x55)    # #555555


def _set_cell_background(cell, fill_hex: str) -> None:
    """Set the background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)


def _set_cell_margins(cell, top: int = 120, bottom: int = 120, left: int = 150, right: int = 150) -> None:
    """Set cell padding (in twentieths of a point / dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def _style_table(table, col_widths=None) -> None:
    """Apply standard corporate borders and column widths."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="6" w:space="0" w:color="123456"/>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="123456"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        f'  <w:insideV w:val="none"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)
    
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = width


def compile_approval_note(payload: ApprovalNotePayload, output_path: Path) -> Path:
    """
    Compile a formal executive Word Approval Note per MRPL corporate standards.
    
    Args:
        payload: Complete data model containing line inspection, formula steps, and recommendations.
        output_path: Destination path for the .docx file.
        
    Returns:
        Path: Resolved output path.
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    doc = Document()
    
    # Page setup: Standard margins (0.75 in)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
    # -------------------------------------------------------------
    # 1. Header / Letterhead
    # -------------------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(2)
    
    run_title = title_p.add_run("MANGALORE REFINERY AND PETROCHEMICALS LIMITED (MRPL)")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(15)
    run_title.font.bold = True
    run_title.font.color.rgb = MRPL_NAVY
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(4)
    run_sub = subtitle_p.add_run("Process Maintenance & Reliability Engineering Division — Technical Evaluation & Approval Note")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(10.5)
    run_sub.font.italic = True
    run_sub.font.color.rgb = MUTED_GRAY
    
    # Divider Rule
    rule_p = doc.add_paragraph()
    rule_p.paragraph_format.space_after = Pt(14)
    rule_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_rule = rule_p.add_run("━" * 58)
    run_rule.font.color.rgb = MRPL_NAVY
    run_rule.font.size = Pt(8)
    
    # -------------------------------------------------------------
    # 2. Document Metadata Table
    # -------------------------------------------------------------
    insp = payload.inspection_data
    meta_table = doc.add_table(rows=2, cols=4)
    _style_table(meta_table, [Inches(1.6), Inches(2.1), Inches(1.5), Inches(1.8)])
    
    meta_data = [
        [("Ref. Docket No:", payload.reference_number), ("Inspection Date:", insp.inspection_date)],
        [("Unit / Area:", payload.unit_name), ("Piping Line Tag:", insp.line_tag)]
    ]
    
    for r_idx, row_pair in enumerate(meta_data):
        row = meta_table.rows[r_idx]
        col_ptr = 0
        for label, val in row_pair:
            # Label cell
            lbl_cell = row.cells[col_ptr]
            _set_cell_background(lbl_cell, "F2F4F8")
            _set_cell_margins(lbl_cell, top=80, bottom=80)
            p = lbl_cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(label)
            r.font.name = "Arial"
            r.font.size = Pt(9)
            r.font.bold = True
            r.font.color.rgb = MRPL_NAVY
            
            # Value cell
            val_cell = row.cells[col_ptr + 1]
            _set_cell_margins(val_cell, top=80, bottom=80)
            p = val_cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(val))
            r.font.name = "Arial"
            r.font.size = Pt(9)
            r.font.color.rgb = TEXT_DARK
            
            col_ptr += 2
            
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    
    # -------------------------------------------------------------
    # 3. Section: Ultrasonic Inspection & Operating Parameters
    # -------------------------------------------------------------
    sec1_p = doc.add_paragraph()
    sec1_p.paragraph_format.space_before = Pt(8)
    sec1_p.paragraph_format.space_after = Pt(4)
    run_sec1 = sec1_p.add_run("1. In-Service Ultrasonic Thickness Inspection (UTG) Summary")
    run_sec1.font.name = "Arial"
    run_sec1.font.size = Pt(11)
    run_sec1.font.bold = True
    run_sec1.font.color.rgb = MRPL_NAVY
    
    utg_table = doc.add_table(rows=6, cols=3)
    _style_table(utg_table, [Inches(3.2), Inches(1.8), Inches(2.0)])
    
    # Header Row
    hdr = utg_table.rows[0]
    for i, title in enumerate(["Inspection Parameter", "Measured / Evaluated Value", "Statutory Standard / Baseline"]):
        cell = hdr.cells[i]
        _set_cell_background(cell, "123456")
        _set_cell_margins(cell, top=100, bottom=100)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(title)
        r.font.name = "Arial"
        r.font.size = Pt(9)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    utg_rows = [
        ("Nominal Design Wall Thickness (t_nominal)", f"{insp.nominal_thickness_mm:.2f} mm", "Original Drawing Spec"),
        ("Actual Measured Wall Thickness (t_actual)", f"{insp.actual_thickness_mm:.2f} mm", "UTG NDT Calibration"),
        ("Minimum Required Wall Thickness (t_min)", f"{insp.minimum_required_thickness_mm:.2f} mm", "ASME B31.3 / API 570"),
        ("Calculated Corrosion Rate (C_r)", f"{insp.corrosion_rate_mm_year:.3f} mm/year", "Operational History"),
        ("Estimated Remaining Safe Service Life", f"{insp.remaining_life_years:.2f} years", "Critical Threshold < 5.0 yrs")
    ]
    
    for r_idx, (param, val, baseline) in enumerate(utg_rows, start=1):
        row = utg_table.rows[r_idx]
        bg = "FFFFFF" if r_idx % 2 != 0 else "F9FBFD"
        for c_idx, cell_text in enumerate([param, val, baseline]):
            cell = row.cells[c_idx]
            _set_cell_background(cell, bg)
            _set_cell_margins(cell, top=80, bottom=80)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(cell_text)
            r.font.name = "Arial"
            r.font.size = Pt(9)
            if c_idx == 1:
                r.font.bold = True
                if "Remaining" in param and insp.remaining_life_years < 5.0:
                    r.font.color.rgb = CRITICAL_RED
                else:
                    r.font.color.rgb = MRPL_NAVY
            else:
                r.font.color.rgb = TEXT_DARK
                
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    
    # -------------------------------------------------------------
    # 4. Section: Mathematical Derivation & Engineering Basis
    # -------------------------------------------------------------
    sec2_p = doc.add_paragraph()
    sec2_p.paragraph_format.space_before = Pt(6)
    sec2_p.paragraph_format.space_after = Pt(4)
    run_sec2 = sec2_p.add_run("2. Engineering Basis & Formula Derivation Steps (API 570 / ASME B31.3)")
    run_sec2.font.name = "Arial"
    run_sec2.font.size = Pt(11)
    run_sec2.font.bold = True
    run_sec2.font.color.rgb = MRPL_NAVY
    
    for step in payload.formula_derivation_steps:
        step_p = doc.add_paragraph(style='List Bullet')
        step_p.paragraph_format.space_after = Pt(3)
        r = step_p.add_run(step)
        r.font.name = "Arial"
        r.font.size = Pt(9)
        r.font.color.rgb = TEXT_DARK
        
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    
    # -------------------------------------------------------------
    # 5. Section: Statutory Action & Recommended Mitigation Box
    # -------------------------------------------------------------
    sec3_p = doc.add_paragraph()
    sec3_p.paragraph_format.space_before = Pt(6)
    sec3_p.paragraph_format.space_after = Pt(4)
    run_sec3 = sec3_p.add_run("3. Statutory Mandate & Operational Recommendation")
    run_sec3.font.name = "Arial"
    run_sec3.font.size = Pt(11)
    run_sec3.font.bold = True
    run_sec3.font.color.rgb = MRPL_NAVY
    
    rec_table = doc.add_table(rows=1, cols=1)
    rec_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rec_cell = rec_table.rows[0].cells[0]
    rec_cell.width = Inches(7.0)
    
    # Highlighted callout box
    is_critical = insp.remaining_life_years < 5.0
    box_bg = "FFF5F5" if is_critical else "F0FDF4"
    border_color = "C00000" if is_critical else "166534"
    _set_cell_background(rec_cell, box_bg)
    _set_cell_margins(rec_cell, top=140, bottom=140, left=180, right=180)
    
    # Border for recommendation box
    tcPr = rec_cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="12" w:space="0" w:color="{border_color}"/>'
        f'  <w:left w:val="single" w:sz="18" w:space="0" w:color="{border_color}"/>'
        f'  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="{border_color}"/>'
        f'  <w:right w:val="single" w:sz="12" w:space="0" w:color="{border_color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p_mandate = rec_cell.paragraphs[0]
    p_mandate.paragraph_format.space_after = Pt(4)
    r_tag = p_mandate.add_run(f"MANDATORY ACTION: {insp.mandatory_action}\n")
    r_tag.font.name = "Arial"
    r_tag.font.size = Pt(10)
    r_tag.font.bold = True
    r_tag.font.color.rgb = CRITICAL_RED if is_critical else MRPL_NAVY
    
    r_rec = p_mandate.add_run(f"Recommendation: {payload.recommended_action}")
    r_rec.font.name = "Arial"
    r_rec.font.size = Pt(9.5)
    r_rec.font.color.rgb = TEXT_DARK
    
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    
    # -------------------------------------------------------------
    # 6. Signatory Approval Block
    # -------------------------------------------------------------
    sig_table = doc.add_table(rows=3, cols=3)
    _style_table(sig_table, [Inches(2.3), Inches(2.3), Inches(2.4)])
    
    roles = [
        ("Inspected & Computed By:", insp.inspector_name, "Non-Destructive Testing (NDT) Lead"),
        ("Reviewed & Verified By:", "Superintendent Engineer (Maintenance)", "Process Reliability Section"),
        ("Final Statutory Approval:", "Chief General Manager (Inspection)", "MRPL Refinery Division")
    ]
    
    for c_idx, (role_label, name, title) in enumerate(roles):
        cell_lbl = sig_table.rows[0].cells[c_idx]
        _set_cell_margins(cell_lbl, top=60, bottom=40)
        p = cell_lbl.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(role_label)
        r.font.name = "Arial"
        r.font.size = Pt(8.5)
        r.font.color.rgb = MUTED_GRAY
        
        # Signature Line cell
        cell_sig = sig_table.rows[1].cells[c_idx]
        _set_cell_margins(cell_sig, top=400, bottom=40)  # blank space for wet signature
        p = cell_sig.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("_______________________________")
        r.font.name = "Arial"
        r.font.size = Pt(9)
        r.font.color.rgb = MUTED_GRAY
        
        # Name and Title cell
        cell_name = sig_table.rows[2].cells[c_idx]
        _set_cell_margins(cell_name, top=40, bottom=60)
        p = cell_name.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r_name = p.add_run(f"{name}\n")
        r_name.font.name = "Arial"
        r_name.font.size = Pt(8.5)
        r_name.font.bold = True
        r_name.font.color.rgb = TEXT_DARK
        
        r_title = p.add_run(title)
        r_title.font.name = "Arial"
        r_title.font.size = Pt(7.5)
        r_title.font.italic = True
        r_title.font.color.rgb = MUTED_GRAY

    # Save to destination
    doc.save(str(output_path))
    return output_path
